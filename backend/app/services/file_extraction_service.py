import os
import io
import time
import hashlib
import tempfile
from typing import Tuple, Optional, Dict, Any, Union
import logging
import mimetypes
from pathlib import Path

# 文档处理相关导入
import PyPDF2
from docx import Document
import aiofiles

# OCR服务导入
from app.services.ocr_service import ocr_service
from app.config import settings

logger = logging.getLogger(__name__)


class FileExtractionCache:
    """文件提取结果缓存"""
    
    def __init__(self, ttl: int = 3600):
        self.cache: Dict[str, Tuple[str, float, Dict]] = {}  # {hash: (text, timestamp, metadata)}
        self.ttl = ttl
    
    def _get_content_hash(self, content: bytes, filename: str) -> str:
        """生成内容哈希值"""
        hasher = hashlib.md5()
        hasher.update(content)
        hasher.update(filename.encode('utf-8'))
        return hasher.hexdigest()
    
    def get(self, content: bytes, filename: str) -> Optional[Tuple[str, Dict]]:
        """获取缓存结果"""
        if not settings.rag_enable_cache:
            return None
            
        content_hash = self._get_content_hash(content, filename)
        if content_hash in self.cache:
            text, timestamp, metadata = self.cache[content_hash]
            if time.time() - timestamp < self.ttl:
                logger.info(f"使用缓存的文件提取结果: {filename}")
                return text, metadata
            else:
                del self.cache[content_hash]
        return None
    
    def set(self, content: bytes, filename: str, text: str, metadata: Dict):
        """设置缓存结果"""
        if not settings.rag_enable_cache:
            return
            
        content_hash = self._get_content_hash(content, filename)
        self.cache[content_hash] = (text, time.time(), metadata)
    
    def clear(self):
        """清空缓存"""
        self.cache.clear()


class FileExtractionService:
    """统一的文件文本提取服务"""
    
    def __init__(self):
        self.cache = FileExtractionCache(ttl=settings.rag_cache_ttl)
        
        # 支持的文件类型映射
        self.supported_types = {
            # PDF文件
            'application/pdf': self._extract_from_pdf,
            
            # 文本文件
            'text/plain': self._extract_from_text,
            
            # Word文档
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'application/msword': self._extract_from_doc,  # 🔧 老版Word格式(.doc)使用专门的处理方法
            
            # 图片文件（需要OCR）
            'image/png': self._extract_from_image,
            'image/jpeg': self._extract_from_image,
            'image/jpg': self._extract_from_image,
            'image/bmp': self._extract_from_image,
            'image/tiff': self._extract_from_image,
            'image/webp': self._extract_from_image,
            
            # 音频文件（占位符，未来可扩展语音识别）
            'audio/wav': self._extract_from_audio,
            'audio/mpeg': self._extract_from_audio,
            'audio/mp3': self._extract_from_audio,
        }
        
        logger.info("文件提取服务初始化完成")
    
    async def extract_text_from_file(
        self, 
        file_content: bytes, 
        filename: str, 
        file_type: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        从文件中提取文本内容
        
        Args:
            file_content: 文件内容字节数据
            filename: 文件名
            file_type: 文件MIME类型（可选，会自动检测）
            
        Returns:
            Tuple[str, Dict]: (提取的文本内容, 元数据信息)
        """
        start_time = time.time()
        
        try:
            # 检查缓存
            cached_result = self.cache.get(file_content, filename)
            if cached_result:
                return cached_result
            
            # 检测文件类型
            if not file_type:
                file_type = self._detect_file_type(file_content, filename)
            
            logger.info(f"开始提取文件文本: {filename}, 类型: {file_type}, 大小: {len(file_content)} bytes")
            
            # 根据文件类型选择提取方法
            if file_type in self.supported_types:
                extractor = self.supported_types[file_type]
                text, extraction_metadata = await extractor(file_content, filename)
            else:
                # 不支持的文件类型，尝试作为文本处理
                logger.warning(f"不支持的文件类型 {file_type}，尝试作为文本处理: {filename}")
                text, extraction_metadata = await self._extract_from_text(file_content, filename)
            
            # 基本清理和验证
            text = self._clean_extracted_text(text)
            
            # 生成完整的元数据
            metadata = {
                'filename': filename,
                'file_type': file_type,
                'file_size': len(file_content),
                'extraction_time': time.time() - start_time,
                'text_length': len(text),
                'word_count': len(text.split()) if text else 0,
                **extraction_metadata
            }
            
            # 缓存结果
            self.cache.set(file_content, filename, text, metadata)
            
            logger.info(f"文件提取完成: {filename}, 文本长度: {len(text)}, 耗时: {metadata['extraction_time']:.2f}秒")
            return text, metadata
            
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"文件提取失败 {filename}: {str(e)}"
            logger.error(error_msg)
            
            # 返回错误信息和基本元数据
            metadata = {
                'filename': filename,
                'file_type': file_type or 'unknown',
                'file_size': len(file_content),
                'extraction_time': processing_time,
                'error': str(e),
                'extraction_method': 'error'
            }
            
            # 返回文件名作为内容，以便后续处理不会完全失败
            fallback_text = f"文档: {filename}\n[文本提取失败: {str(e)}]"
            return fallback_text, metadata
    
    def _detect_file_type(self, file_content: bytes, filename: str) -> str:
        """检测文件类型"""
        try:
            # 先尝试根据文件名检测
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type:
                return mime_type
            
            # 检查文件签名（魔数）
            if file_content.startswith(b'%PDF'):
                return 'application/pdf'
            elif file_content.startswith(b'PK\x03\x04'):
                # 检查是否是DOCX文件
                if b'word/' in file_content[:1024]:
                    return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_content.startswith((b'\xff\xd8\xff', b'\x89PNG')):
                # JPEG或PNG图片
                return 'image/jpeg' if file_content.startswith(b'\xff\xd8\xff') else 'image/png'
            
            # 默认尝试作为文本
            return 'text/plain'
            
        except Exception as e:
            logger.warning(f"文件类型检测失败 {filename}: {e}")
            return 'application/octet-stream'
    
    async def _extract_from_pdf(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从PDF文件提取文本"""
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                # 首先检测PDF类型（文本PDF vs 扫描PDF）
                is_text_pdf, extracted_text, char_count = await ocr_service.detect_pdf_text_content(temp_path)
                
                if is_text_pdf and char_count > 50:
                    # 文本PDF，直接提取文本
                    logger.info(f"检测到文本PDF: {filename}, 字符数: {char_count}")
                    
                    # 使用更完整的文本提取
                    full_text = await ocr_service.extract_full_pdf_text(temp_path)
                    
                    metadata = {
                        'extraction_method': 'text_pdf',
                        'is_text_pdf': True,
                        'is_pdf': True,
                        'char_count': len(full_text),
                        'detected_char_count': char_count,  # 检测阶段的字符数
                        'confidence': 1.0,  # 文本PDF置信度为1.0
                        'processing_status': 'text_pdf_extracted'
                    }
                    
                    return full_text, metadata
                    
                else:
                    # 扫描PDF，使用OCR
                    logger.info(f"检测到扫描PDF: {filename}, 使用OCR处理")
                    
                    text, confidence, processing_time = await ocr_service.extract_text_from_pdf(temp_path)
                    
                    metadata = {
                        'extraction_method': 'ocr_pdf',
                        'is_text_pdf': False,
                        'is_pdf': True,
                        'char_count': len(text),
                        'detected_char_count': char_count,  # 检测阶段的字符数
                        'confidence': confidence,
                        'ocr_processing_time': processing_time,
                        'processing_status': f'scanned_pdf_ocr_completed'
                    }
                    
                    return text, metadata
                    
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"PDF文本提取失败: {e}")
            raise Exception(f"PDF文本提取失败: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从DOCX文件提取文本"""
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                doc = Document(temp_path)
                text_parts = []
                
                # 提取段落文本
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # 提取表格文本
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(' | '.join(row_text))
                
                text = '\n'.join(text_parts)
                
                metadata = {
                    'extraction_method': 'docx',
                    'is_pdf': False,
                    'is_text_pdf': False,
                    'char_count': len(text),
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'confidence': 1.0,  # DOCX文本提取置信度高
                    'processing_status': 'docx_extracted'
                }
                
                return text, metadata
                
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"DOCX文本提取失败: {e}")
            raise Exception(f"DOCX文本提取失败: {str(e)}")
    
    async def _extract_from_doc(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从老版Word文档(.doc)提取文本"""
        try:
            logger.info(f"处理老版Word文档(.doc): {filename}")
            
            # 对于.doc文件，我们有几种选择：
            # 1. 使用python-docx2txt库（如果安装了）
            # 2. 使用antiword工具（如果系统有）
            # 3. 尝试用docx库处理（可能失败）
            # 4. 回退到OCR处理
            
            # 首先尝试导入docx2txt库
            try:
                import docx2txt
                # 创建临时文件
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(file_content)
                    temp_path = temp_file.name
                
                try:
                    # 使用docx2txt提取文本
                    text = docx2txt.process(temp_path)
                    
                    if text and text.strip():
                        metadata = {
                            'extraction_method': 'docx2txt',
                            'is_pdf': False,
                            'is_text_pdf': False,
                            'char_count': len(text),
                            'confidence': 0.9,  # docx2txt对.doc文件支持较好
                            'processing_status': 'doc_extracted_with_docx2txt'
                        }
                        
                        logger.info(f"使用docx2txt成功提取.doc文档: {filename}, 字符数: {len(text)}")
                        return text.strip(), metadata
                    
                finally:
                    # 清理临时文件
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                        
            except ImportError:
                logger.warning("docx2txt库未安装，尝试其他方法处理.doc文件")
            
            # 尝试使用antiword命令行工具（如果系统有安装）
            try:
                import subprocess
                
                # 创建临时文件
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(file_content)
                    temp_path = temp_file.name
                
                try:
                    # 运行antiword命令
                    result = subprocess.run(
                        ['antiword', temp_path], 
                        capture_output=True, 
                        text=True, 
                        timeout=30
                    )
                    
                    if result.returncode == 0 and result.stdout.strip():
                        text = result.stdout.strip()
                        
                        metadata = {
                            'extraction_method': 'antiword',
                            'is_pdf': False,
                            'is_text_pdf': False,
                            'char_count': len(text),
                            'confidence': 0.95,  # antiword对.doc文件支持很好
                            'processing_status': 'doc_extracted_with_antiword'
                        }
                        
                        logger.info(f"使用antiword成功提取.doc文档: {filename}, 字符数: {len(text)}")
                        return text, metadata
                        
                finally:
                    # 清理临时文件
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)
                        
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
                logger.warning(f"antiword处理失败: {e}")
            
            # 如果上述方法都失败，尝试用docx库处理（可能会失败，但值得一试）
            try:
                logger.info(f"尝试用docx库处理.doc文件: {filename}")
                return await self._extract_from_docx(file_content, filename)
            except Exception as docx_error:
                logger.warning(f"docx库无法处理.doc文件: {docx_error}")
            
            # 最后回退：将.doc文件当作二进制文件，提供基本信息
            logger.warning(f"无法提取.doc文件内容，返回基本信息: {filename}")
            
            fallback_text = f"""文档: {filename}
文件类型: Microsoft Word 97-2003 文档 (.doc)
文件大小: {len(file_content)} 字节

⚠️ 文本提取说明:
老版Word文档(.doc)需要专门的工具进行文本提取。
建议:
1. 将文档转换为.docx格式后重新上传
2. 安装docx2txt库: pip install docx2txt
3. 安装antiword工具: apt-get install antiword (Linux)

此文档已保存，可以稍后处理。"""
            
            metadata = {
                'extraction_method': 'doc_fallback',
                'is_pdf': False,
                'is_text_pdf': False,
                'char_count': len(fallback_text),
                'confidence': 0.0,  # 未提取到真实内容
                'processing_status': 'doc_needs_conversion',
                'note': '需要安装额外工具或转换格式'
            }
            
            return fallback_text, metadata
            
        except Exception as e:
            logger.error(f"DOC文本提取失败: {e}")
            # 返回错误信息而不是抛出异常，避免整个流程中断
            error_text = f"文档: {filename}\n[.doc文件文本提取失败: {str(e)}]\n\n建议将文档转换为.docx格式后重新上传。"
            
            metadata = {
                'extraction_method': 'doc_error',
                'is_pdf': False,
                'is_text_pdf': False,
                'char_count': len(error_text),
                'confidence': 0.0,
                'processing_status': 'doc_extraction_failed',
                'error': str(e)
            }
            
            return error_text, metadata
    
    async def _extract_from_text(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从文本文件提取内容"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'utf-16', 'ascii']
            
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # 如果所有编码都失败，使用错误忽略模式
                text = file_content.decode('utf-8', errors='ignore')
                used_encoding = 'utf-8 (with errors ignored)'
            
            metadata = {
                'extraction_method': 'text',
                'is_pdf': False,
                'is_text_pdf': False,
                'char_count': len(text),
                'encoding': used_encoding,
                'confidence': 1.0 if '(with errors ignored)' not in used_encoding else 0.8,
                'processing_status': 'text_extracted'
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            raise Exception(f"文本提取失败: {str(e)}")
    
    async def _extract_from_image(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从图片文件提取文本（OCR）"""
        try:
            # 创建临时文件
            file_ext = Path(filename).suffix.lower()
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                # 使用OCR服务提取文本
                text, confidence, processing_time = await ocr_service.extract_text_from_image(temp_path)
                
                metadata = {
                    'extraction_method': 'ocr_image',
                    'is_pdf': False,
                    'is_text_pdf': False,
                    'char_count': len(text),
                    'confidence': confidence,
                    'ocr_processing_time': processing_time,
                    'processing_status': 'image_ocr_completed'
                }
                
                return text, metadata
                
            finally:
                # 清理临时文件
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except Exception as e:
            logger.error(f"图片OCR提取失败: {e}")
            raise Exception(f"图片OCR提取失败: {str(e)}")
    
    async def _extract_from_audio(self, file_content: bytes, filename: str) -> Tuple[str, Dict]:
        """从音频文件提取文本（占位符，未来可扩展语音识别）"""
        # 目前音频文件不进行文本提取，返回基本信息
        logger.info(f"音频文件暂不支持文本提取: {filename}")
        
        # 检测音频格式
        audio_format = "unknown"
        if filename.lower().endswith('.wav'):
            audio_format = "wav"
        elif filename.lower().endswith(('.mp3', '.mpeg')):
            audio_format = "mp3"
        
        metadata = {
            'extraction_method': 'audio_placeholder',
            'is_pdf': False,
            'is_text_pdf': False,
            'char_count': 0,
            'audio_format': audio_format,
            'confidence': 0.0,  # 音频未处理，置信度为0
            'processing_status': 'audio_file_uploaded',
            'note': '音频文件已上传，暂不支持自动语音识别转文本'
        }
        
        # 返回描述性文本
        text = f"音频文件: {filename}\n[格式: {audio_format.upper()}]\n[暂不支持自动语音识别转文本]"
        
        return text, metadata
    
    def _clean_extracted_text(self, text: str) -> str:
        """清理提取的文本"""
        if not text:
            return ""
        
        # 基本清理
        text = text.strip()
        
        # 移除过多的空行
        import re
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # 移除行首行尾多余空格
        lines = []
        for line in text.split('\n'):
            lines.append(line.strip())
        
        return '\n'.join(lines)
    
    def get_supported_file_types(self) -> Dict[str, str]:
        """获取支持的文件类型"""
        return {
            'application/pdf': 'PDF文档',
            'text/plain': '文本文件',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'Word文档 (.docx)',
            'application/msword': 'Word文档 (.doc)',
            'image/png': 'PNG图片',
            'image/jpeg': 'JPEG图片',
            'image/jpg': 'JPG图片',
            'image/bmp': 'BMP图片',
            'image/tiff': 'TIFF图片',
            'image/webp': 'WebP图片',
            'audio/wav': 'WAV音频文件',
            'audio/mpeg': 'MP3音频文件',
            'audio/mp3': 'MP3音频文件'
        }
    
    def is_supported_file_type(self, file_type: str) -> bool:
        """检查文件类型是否支持"""
        return file_type in self.supported_types
    
    def clear_cache(self):
        """清空缓存"""
        self.cache.clear()
        logger.info("文件提取缓存已清空")


# 创建全局服务实例
file_extraction_service = FileExtractionService()
